"""Pipeline orchestrator for project analysis."""
import requests
from typing import Any, Callable
from docx.oxml.ns import qn
from docx.shared import RGBColor
import json
import subprocess
from pathlib import Path

DEFAULT_REPORT_DIR = Path(r"E:\work\word\study\word")
DEFAULT_ENV_PATH = Path(r"C:\Users\Luo\AppData\Local\hermes\.env")
DEFAULT_INDEX_PATH = Path(r"E:\work\word\study\word\项目分析索引.md")
DEFAULT_DOWNLOAD_ROOT = Path(r"E:\work\word\study\他人项目")
DEFAULT_RULER_PATH = Path(r"E:\work\word\study\他人项目\ai-best-practices\pareto-six-dimensions.md")
DEFAULT_CLASSIFICATION_PATH = Path(
    r"C:\Users\Luo\AppData\Local\hermes\profiles\learning\skills\thinking-partner\six-dimensions-analyzer\references\classification.md"
)
DEFAULT_UID = "383325863"
DEFAULT_SPACE_URL = "https://space.bilibili.com/383325863/video"
DEFAULT_USER_PROFILE = "AI学习阶段；关注Agent架构和工程实践。"


def _fail(stage: str, repo_key: str | None, error: dict[str, Any]) -> dict[str, Any]:
    return {
        "status": "failed",
        "failed_stage": stage,
        "error_code": error.get("error_code", f"{stage.upper()}_FAILED"),
        "message": error.get("message", ""),
        "recovery": "continue_next_project",
        "repo_key": repo_key,
    }


def _mark(state: dict[str, Any], status: str) -> None:
    state["status"] = status
    state.setdefault("transitions", []).append(status)


def load_env(path: Path | str = DEFAULT_ENV_PATH) -> dict[str, str]:
    env_path = Path(path)
    values: dict[str, str] = {}
    if not env_path.exists():
        return values
    for raw_line in env_path.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        key, value = line.split("=", 1)
        values[key.strip()] = value.strip().strip('"').strip("'")
    return values


def build_deepseek_client(env: dict[str, str]) -> Callable[[str], str]:
    api_key = env.get("TOKENESS_API_KEY") or env.get("DEEPSEEK_API_KEY")
    model = env.get("TOKENESS_MODEL") or env.get("DEEPSEEK_MODEL", "deepseek-chat")
    base_url = env.get("TOKENESS_BASE_URL") or env.get("DEEPSEEK_BASE_URL", "https://api.deepseek.com/chat/completions")
    if not api_key:
        raise RuntimeError("TOKENESS_API_KEY or DEEPSEEK_API_KEY is missing")

    def client(prompt: str) -> str:
        payload = {"model": model, "messages": [{"role": "user", "content": prompt}], "temperature": 0.2}
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        resp = requests.post(base_url, json=payload, headers=headers, timeout=120,
                            proxies={"https": "http://127.0.0.1:7897"})
        resp.raise_for_status()
        return resp.json()["choices"][0]["message"]["content"]

    return client


def _load_text_if_exists(path: Path | str) -> str:
    candidate = Path(path)
    if not candidate.exists():
        return ""
    return candidate.read_text(encoding="utf-8", errors="replace")


def _search_projects(**kwargs: Any) -> dict[str, Any]:
    import searcher

    return searcher.search_projects(**kwargs)


def _scan_project(path: Path | str, repo_key: str) -> dict[str, Any]:
    import scanner

    return scanner.scan_project(path, repo_key)


def _query_index(path: Path | str, repo_key: str) -> dict[str, Any]:
    import indexer

    return indexer.query_index(path, repo_key)


def _update_index(path: Path | str, entry: dict[str, Any]) -> dict[str, Any]:
    import indexer

    return indexer.update_index(path, entry)


def _analyze_profile(profile: dict[str, Any], client: Callable[[str], str]) -> dict[str, Any]:
    import analyzer_perspective1

    return analyzer_perspective1.analyze_profile(profile, client)


def _analyze_architecture(profile: dict[str, Any], client: Callable[[str], str]) -> dict[str, Any]:
    import analyzer_perspective2

    return analyzer_perspective2.analyze_architecture(profile, client)


def _analyze_engineering(
    profile: dict[str, Any],
    client: Callable[[str], str],
    ruler_text: str = "",
    classification_text: str = "",
) -> dict[str, Any]:
    import analyzer_perspective3

    return analyzer_perspective3.analyze_engineering(profile, client, ruler_text, classification_text)


def _analyze_hidden_risks(profile: dict[str, Any], client: Callable[[str], str]) -> dict[str, Any]:
    import analyzer_perspective4b

    return analyzer_perspective4b.analyze_hidden_risks(profile, client)


def run_json_subprocess(
    command: list[str],
    payload: dict[str, Any],
    timeout_seconds: int = 120,
) -> dict[str, Any]:
    try:
        completed = subprocess.run(
            command,
            input=json.dumps(payload, ensure_ascii=False),
            text=True,
            capture_output=True,
            timeout=timeout_seconds,
            check=False,
        )
    except subprocess.TimeoutExpired as error:
        return {
            "schema_version": 1,
            "ok": False,
            "stage": "subprocess",
            "error_code": "SUBPROCESS_TIMEOUT",
            "message": str(error),
            "recoverable": True,
        }

    if completed.returncode != 0:
        return {
            "schema_version": 1,
            "ok": False,
            "stage": "subprocess",
            "error_code": "SUBPROCESS_FAILED",
            "message": completed.stderr.strip(),
            "recoverable": True,
        }

    try:
        return json.loads(completed.stdout)
    except json.JSONDecodeError as error:
        return {
            "schema_version": 1,
            "ok": False,
            "stage": "subprocess",
            "error_code": "SUBPROCESS_JSON_INVALID",
            "message": str(error),
            "recoverable": True,
        }


def _analyze_dialectic(results: dict[str, dict[str, Any]], client: Callable[[str], str]) -> dict[str, Any]:
    import analyzer_perspective4a

    return analyzer_perspective4a.analyze_dialectic(results, client)


def _score_relevance(
    results: dict[str, dict[str, Any]],
    user_profile: str,
    client: Callable[[str], str],
) -> dict[str, Any]:
    import analyzer_perspective5

    return analyzer_perspective5.score_relevance(results, user_profile, client)


CN_NUMS = ["零", "一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]


def _next_batch_num(report_dir: Path | str) -> int:
    """从已有报告文件名中提取最大批次号，返回下一个批次号。"""
    import re
    directory = Path(report_dir)
    if not directory.exists():
        return 1
    max_batch = 0
    cn_to_num = {cn: i for i, cn in enumerate(CN_NUMS)}
    for f in directory.glob("*.docx"):
        m = re.match(r"^([一二三四五六七八九十]+)-", f.name)
        if m:
            cn = m.group(1)
            num = 0
            if len(cn) == 1:
                num = cn_to_num.get(cn, 0)
            elif cn == "十":
                num = 10
            max_batch = max(max_batch, num)
    return max_batch + 1


def _make_report_name(repo_key: str, date_str: str, batch_num: int, seq: int, source: str = "GitHub热榜") -> str:
    """生成报告文件名：一-01-owner-repo-分析报告-20260720-GitHub热榜.docx"""
    safe_key = repo_key.replace("/", "-").replace("\\", "-")
    batch_cn = CN_NUMS[batch_num] if batch_num < len(CN_NUMS) else str(batch_num)
    return f"{batch_cn}-{seq:02d}-{safe_key}-分析报告-{date_str}-{source}.docx"


def render_markdown_report(item: dict[str, Any]) -> str:
    repo_key = item.get("repo_key", "")
    score = item.get("relevance_score", 0)
    analysis = item.get("analysis", {})

    sections = [
        f"# {repo_key} — 五视角分析",
        f"**相关性**：{score}",
        "## 1. 它是什么",
        str(analysis.get("perspective1", {}).get("content", "")),
        "## 2. 架构骨架",
        str(analysis.get("perspective2", {}).get("content", "")),
        "## 3. 工程评估",
        str(analysis.get("perspective3", {}).get("content", "")),
        "## 4A. 矛盾分析",
        str(analysis.get("perspective4a", {}).get("content", "")),
        "## 4B. 裂隙分析",
        str(analysis.get("perspective4b", {}).get("content", "")),
        "## 5. 适用性分析",
        str(analysis.get("perspective5", {}).get("content", "")),
    ]
    return "\n\n".join(sections) + "\n"


def write_markdown_report(item: dict[str, Any], report_dir: Path | str = DEFAULT_REPORT_DIR) -> Path:
    directory = Path(report_dir)
    directory.mkdir(parents=True, exist_ok=True)
    report_path = directory / f"{_safe_report_name(str(item.get('repo_key', 'project')))}-分析报告.md"
    temp_path = report_path.with_suffix(report_path.suffix + ".tmp")
    temp_path.write_text(render_markdown_report(item), encoding="utf-8")
    temp_path.replace(report_path)
    return report_path


def write_word_report(item: dict[str, Any], report_dir: Path | str = DEFAULT_REPORT_DIR,
                      date_str: str = "", batch_num: int = 1, seq: int = 1) -> Path:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import re

    repo_key = item.get("repo_key", "")
    score = item.get("relevance_score", 0)
    analysis = item.get("analysis", {})

    doc = Document()
    # Set default font
    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.3
    # Set East-Asian font
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:eastAsia"), "楷体")
    rPr.insert(0, rFonts)

    # Title
    h = doc.add_heading(f"{repo_key} — 五视角分析", level=0)
    for run in h.runs:
        _set_run_font(run, "黑体", size=Pt(18), bold=True)
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()
    r = p.add_run(f"相关性：{score}")
    _set_run_font(r, "楷体", size=Pt(12))

    sections = [
        ("1. 它是什么", "perspective1"),
        ("2. 架构骨架", "perspective2"),
        ("3. 工程评估", "perspective3"),
        ("4. 矛盾分析", "perspective4a"),
        ("5. 裂隙分析", "perspective4b"),
        ("6. 适用性分析", "perspective5"),
    ]
    section_map = dict(sections)

    first = True
    for title_text, key in sections:
        if not first:
            doc.add_page_break()
        first = False

        h = doc.add_heading(title_text, level=1)
        for run in h.runs:
            _set_run_font(run, "黑体", size=Pt(16), bold=True)
            run.font.color.rgb = RGBColor(0, 0, 0)

        content = str(analysis.get(key, {}).get("content", ""))
        # Strip code blocks entirely
        content = re.sub(r'```[^`]*```', '', content, flags=re.DOTALL)
        # Clean markdown artifacts
        content = re.sub(r'`[^`]+/[^`]+`', '', content)
        content = re.sub(r'`([^`/ ]+\.[a-z]+)`', r'\1', content)
        content = re.sub(r'\n-{3,}\n', '\n', content)
        # Strip residual triple backticks
        content = content.replace('```', '')
        # Backtick quotes → Chinese quotes (only inline `...`, not multi-line)
        parts = content.split('`')
        result = []
        for pi, p in enumerate(parts):
            if pi % 2 == 1:
                result.append('\u201c' + p + '\u201d')
            else:
                result.append(p)
        content = ''.join(result)
        _write_content(doc, content)

    directory = Path(report_dir)
    directory.mkdir(parents=True, exist_ok=True)
    report_path = directory / _make_report_name(str(repo_key), date_str, batch_num, seq)
    doc.save(str(report_path))
    return report_path


def _set_run_font(run, cn_font, size=None, bold=None, en_font="Times New Roman"):
    """Set both Chinese and English font on a run."""
    run.font.name = en_font
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)
    # Remove existing rFonts if any
    for old in rPr.findall(qn("w:rFonts")):
        rPr.remove(old)
    rPr.insert(0, rFonts)
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold


def _write_content(doc, text):
    """Write content to docx paragraph by paragraph, handling tables inline."""
    from docx.shared import Pt
    import re

    paragraphs = text.split("\n")
    i = 0
    pending = []  # buffer for markdown bold spans within one paragraph

    while i < len(paragraphs):
        line = paragraphs[i].strip()

        # Detect pipe table: lines starting with | and containing |---|
        if line.startswith("|") and "|" in line:
            rows = [line]
            j = i + 1
            while j < len(paragraphs):
                nl = paragraphs[j].strip()
                if nl.startswith("|"):
                    rows.append(nl)
                    j += 1
                else:
                    break
            # Filter out separator rows like |---|---|
            real_rows = [r for r in rows if not re.match(r'^\|[\s\-:|]+\|$', r)]
            if len(real_rows) >= 1:
                _write_table(doc, real_rows)
            i = j
            continue

        if not line:
            i += 1
            continue

        # Skip lines that are just markdown headings (#, ##, ###, ####)
        if re.match(r'^#{1,4}\s', line):
            hcount = len(re.match(r'^(#+)', line).group(1))
            level = min(hcount, 4)
            h = doc.add_heading(re.sub(r'^#{1,4}\s+', '', line), level=level)
            sizes = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(12)}
            for run in h.runs:
                _set_run_font(run, "黑体", size=sizes[level], bold=True)
                run.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        # Handle bold (**text**) and bullet (- or * text)
        line = re.sub(r'^\s*[-*]\s+', '', line)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        # Split by ** for bold spans
        parts = line.split("**")
        for pi, part in enumerate(parts):
            if not part:
                continue
            r = p.add_run(part)
            _set_run_font(r, "楷体", size=Pt(12))
            if pi % 2 == 1:  # odd parts are bold
                r.font.bold = True

        i += 1


def _write_table(doc, rows):
    """Write a list of |col1|col2|...| strings as a Word table."""
    from docx.shared import Pt

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    ncols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=ncols)
    table.style = "Table Grid"

    for ri, row in enumerate(parsed):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = table.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            # Handle **bold** in cell text
            parts = cell_text.split("**")
            for pi, part in enumerate(parts):
                if not part:
                    continue
                r = p.add_run(part)
                _set_run_font(r, "楷体", size=Pt(9))
                if ri == 0 or pi % 2 == 1:
                    r.font.bold = True


def _batch_state_path(report_dir: Path | str, batch_num: int, date_str: str) -> Path:
    return Path(report_dir) / f".batch-{batch_num:02d}-{date_str}.json"


def _load_batch_state(report_dir: Path | str, date_str: str) -> dict | None:
    """Find incomplete batch state for the given date. Returns None if no pending batch."""
    import json as _json
    directory = Path(report_dir)
    for f in sorted(directory.glob(".batch-*-*.json")):
        try:
            state = _json.loads(f.read_text(encoding="utf-8"))
        except Exception:
            continue
        if not state.get("completed") and state.get("pending"):
            return {**state, "_path": str(f)}
    return None


def _save_batch_state(
    report_dir: Path | str, batch_num: int, date_str: str,
    results: list[dict], pending: list[str],
) -> Path:
    import json as _json
    path = _batch_state_path(report_dir, batch_num, date_str)
    path.write_text(_json.dumps({
        "batch_num": batch_num, "date_str": date_str,
        "results": results, "pending": pending,
        "completed": len(pending) == 0,
    }, ensure_ascii=False, indent=2), encoding="utf-8")
    return path


def _merge_and_renumber(
    report_dir: Path | str, batch_num: int, date_str: str,
    prev_results: list[dict], new_results: list[dict],
) -> list[dict]:
    """Merge previous and new results, re-sort by relevance, re-number all files."""
    directory = Path(report_dir)
    batch_cn = CN_NUMS[batch_num] if batch_num < len(CN_NUMS) else str(batch_num)
    pattern = f"{batch_cn}-*-分析报告-{date_str}-GitHub热榜.docx"

    # Remove old batch files
    for old in directory.glob(pattern):
        old.unlink(missing_ok=True)

    # Merge and sort
    all_results = prev_results + new_results
    all_results.sort(key=lambda item: item.get("relevance_score", 0), reverse=True)

    # Write all with correct sequence
    for seq, item in enumerate(all_results, 1):
        write_word_report(item, directory, date_str, batch_num, seq)

    return all_results


def run_cli(
    pipeline: Callable[[], dict[str, Any]],
    input_func: Callable[[str], str] = input,
    print_func: Callable[[str], None] = print,
    report_dir: Path | str = DEFAULT_REPORT_DIR,
    batch_num: int = 1,
) -> dict[str, Any]:
    from datetime import date

    date_str = date.today().strftime("%Y%m%d")
    pipeline_result = pipeline()
    if pipeline_result.get("ok") is not True:
        print_func(f"运行失败：{pipeline_result.get('error_code', 'UNKNOWN')}")
        return pipeline_result

    results = sorted(
        pipeline_result.get("results", []),
        key=lambda value: value.get("relevance_score", 0),
        reverse=True,
    )
    skipped = pipeline_result.get("skipped", [])
    if skipped:
        print_func(f"跳过 {len(skipped)} 个已分析项目：")
        for item in skipped:
            extra = f"（{item['date']}）" if item.get("date") else ""
            print_func(f"  {item['repo_key']} — {item['reason']}{extra}")

    for seq, item in enumerate(results, 1):
        repo_key = item.get("repo_key", "")
        score = item.get("relevance_score", 0)
        overview = item.get("analysis", {}).get("perspective1", {}).get("content", "")
        print_func(f"{repo_key} | 相关性 {score}\n{overview}")
        try:
            action = input_func("回车展开 / n下一个 / q退出：").strip().lower()
        except (EOFError, OSError):
            action = ""  # non-interactive mode: auto-expand all
        if action == "q":
            break
        if action == "n":
            continue
        word_path = write_word_report(item, report_dir, date_str, batch_num, seq)
        print_func(f"报告已写入：{word_path}")
        print_func(render_markdown_report(item))
        print_func(f"报告已写入：{word_path}")

    # Track scan-failed repos for later retry
    pending = []
    for repo_key, st in pipeline_result.get("states", {}).items():
        if st.get("failed_stage") == "scan" and repo_key not in {
            r.get("repo_key") for r in results
        }:
            pending.append(repo_key)

    if pending:
        _save_batch_state(report_dir, batch_num, date_str, results, pending)

    return {"schema_version": 1, "ok": True, "stage": "cli", "pending_retry": pending}


def run_pipeline(
    search_projects: Callable[[], dict[str, Any]],
    query_index: Callable[[str, dict[str, Any] | None], dict[str, Any]],
    scan_project: Callable[[dict[str, Any]], dict[str, Any]],
    run_parallel: Callable[[dict[str, Any]], dict[str, dict[str, Any]]],
    run_dialectic: Callable[[dict[str, dict[str, Any]]], dict[str, Any]],
    run_relevance: Callable[[dict[str, dict[str, Any]]], dict[str, Any]],
    update_index: Callable[[dict[str, Any]], dict[str, Any]],
) -> dict[str, Any]:
    search_result = search_projects()
    if search_result.get("ok") is not True:
        return {
            "schema_version": 1,
            "ok": False,
            "stage": "orchestrator",
            "error_code": search_result.get("error_code", "SEARCH_FAILED"),
            "message": search_result.get("message", ""),
            "recoverable": True,
        }

    states: dict[str, dict[str, Any]] = {}
    results = []
    skipped = []

    for project in search_result.get("projects", []):
        repo_key = project.get("repo_key")
        state = {"repo_key": repo_key, "transitions": []}
        states[repo_key] = state

        first_history = query_index(repo_key, None)
        _mark(state, "history_checked")
        if first_history.get("found"):
            entry = first_history.get("entry", {})
            skipped.append({"repo_key": repo_key, "reason": "已分析过", "date": entry.get("analysis_date", ""), "report": entry.get("report_path", "")})
            state.update(
                {
                    "status": "skipped_analyzed",
                    "matched_repo_keys": [entry.get("repo_key", repo_key)],
                    "same_kind_rule": "exact_repo_key",
                    "skip_reason": "history_match_before_scan",
                }
            )
            continue

        profile = scan_project(project)
        if profile.get("ok") is not True:
            state.update(_fail("scan", repo_key, profile))
            continue
        _mark(state, "scanned")

        second_history = query_index(repo_key, profile)
        if second_history.get("found"):
            entry = second_history.get("entry", {})
            skipped.append({"repo_key": repo_key, "reason": "扫描后匹配已有项目", "date": entry.get("analysis_date", ""), "report": entry.get("report_path", "")})
            state.update(
                {
                    "status": "skipped_analyzed",
                    "matched_repo_keys": [second_history.get("entry", {}).get("repo_key", repo_key)],
                    "same_kind_rule": "profile_similarity",
                    "skip_reason": "history_match_after_scan",
                }
            )
            continue

        parallel_results = run_parallel(profile)
        failed_parallel = next((item for item in parallel_results.values() if item.get("ok") is not True), None)
        if failed_parallel:
            state.update(_fail(failed_parallel.get("stage", "parallel"), repo_key, failed_parallel))
            continue
        _mark(state, "parallel_analyzed")

        dialectic = run_dialectic(parallel_results)
        if dialectic.get("ok") is not True:
            state.update(_fail("perspective4a", repo_key, dialectic))
            continue
        all_results = {**parallel_results, "perspective4a": dialectic}
        _mark(state, "dialectic_analyzed")

        relevance = run_relevance(all_results)
        if relevance.get("ok") is not True:
            state.update(_fail("perspective5", repo_key, relevance))
            continue
        all_results["perspective5"] = relevance
        _mark(state, "fitness_analyzed")

        update_request = {
            "repo_key": repo_key,
            "project_name": repo_key.split("/")[-1] if isinstance(repo_key, str) else "",
            "project_type": profile.get("project_type", ""),
            "tags": profile.get("tags", []),
            "analysis_date": "",
            "relevance": relevance.get("relevance_score", 0),
            "report_path": "",
        }
        update_result = update_index(update_request)
        if update_result.get("ok") is not True:
            state.update(_fail("indexer_update", repo_key, update_result))
            continue
        _mark(state, "indexed")

        results.append(
            {
                "repo_key": repo_key,
                "relevance_score": relevance.get("relevance_score", 0),
                "analysis": all_results,
            }
        )

    results.sort(key=lambda item: item.get("relevance_score", 0), reverse=True)
    return {"schema_version": 1, "ok": True, "stage": "orchestrator", "states": states, "results": results, "skipped": skipped}


def main() -> dict[str, Any]:
    from datetime import date

    env = load_env(DEFAULT_ENV_PATH)
    client = build_deepseek_client(env)
    user_profile = env.get("USER_PROFILE") or DEFAULT_USER_PROFILE
    ruler_text = _load_text_if_exists(DEFAULT_RULER_PATH)
    classification_text = _load_text_if_exists(DEFAULT_CLASSIFICATION_PATH)
    today_str = date.today().strftime("%Y%m%d")

    # ── Retry: check for incomplete batch ──
    batch_state = _load_batch_state(DEFAULT_REPORT_DIR, today_str)
    if batch_state:
        prev_results = batch_state["results"]
        pending_repos = batch_state["pending"]
        batch_num = batch_state["batch_num"]
        print(f"[retry] 批次 {CN_NUMS[batch_num]} 有 {len(pending_repos)} 个项目待补跑")

        new_results = []
        for repo_key in pending_repos:
            owner, repo = repo_key.split("/", 1)
            local_path = DEFAULT_DOWNLOAD_ROOT / f"{owner}-{repo}"
            profile = _scan_project(str(local_path), repo_key)
            if profile.get("ok") is not True:
                print(f"  ✗ {repo_key} 扫描失败，跳过")
                continue

            # Run parallel perspective analysis
            parallel = {
                "perspective1": _analyze_profile(profile, client),
                "perspective2": _analyze_architecture(profile, client),
                "perspective3": _analyze_engineering(profile, client, ruler_text, classification_text),
                "perspective4b": _analyze_hidden_risks(profile, client),
            }
            if any(v.get("ok") is not True for v in parallel.values()):
                print(f"  ✗ {repo_key} 并行分析失败，跳过")
                continue

            dialectic = _analyze_dialectic(parallel, client)
            if dialectic.get("ok") is not True:
                print(f"  ✗ {repo_key} 辩证分析失败，跳过")
                continue

            relevance = _score_relevance({**parallel, "perspective4a": dialectic}, user_profile, client)
            if relevance.get("ok") is not True:
                print(f"  ✗ {repo_key} 适用性分析失败，跳过")
                continue

            new_results.append({
                "repo_key": repo_key,
                "relevance_score": relevance.get("relevance_score", 0),
                "analysis": {**parallel, "perspective4a": dialectic, "perspective5": relevance},
            })
            # Update index
            _update_index(DEFAULT_INDEX_PATH, {
                "repo_key": repo_key,
                "project_name": repo,
                "project_type": profile.get("project_type", ""),
                "tags": profile.get("tags", []),
                "analysis_date": "",
                "relevance": relevance.get("relevance_score", 0),
                "report_path": "",
            })
            print(f"  ✓ {repo_key} 分析完成")

        if new_results:
            _merge_and_renumber(DEFAULT_REPORT_DIR, batch_num, today_str, prev_results, new_results)
            print(f"[retry] 合并完成，共 {len(prev_results) + len(new_results)} 个项目已重排序")
            # Mark batch completed
            _save_batch_state(DEFAULT_REPORT_DIR, batch_num, today_str,
                            prev_results + new_results, [])

        return {"schema_version": 1, "ok": True, "stage": "retry"}

    def pipeline() -> dict[str, Any]:
        def search_projects() -> dict[str, Any]:
            result = _search_projects(download_root=DEFAULT_DOWNLOAD_ROOT)
            if result.get("ok") is not True:
                return result
            repositories = result.get("content", {}).get("repositories", [])
            projects = [
                {
                    "repo_key": item.get("repo_key"),
                    "url": item.get("url"),
                    "local_path": item.get("local_path") or item.get("path"),
                }
                for item in repositories
                if item.get("repo_key")
            ]
            return {**result, "projects": projects}

        return run_pipeline(
            search_projects=search_projects,
            query_index=lambda repo_key, profile=None: _query_index(DEFAULT_INDEX_PATH, repo_key),
            scan_project=lambda project: _scan_project(project.get("local_path", ""), project.get("repo_key", "")),
            run_parallel=lambda profile: {
                "perspective1": _analyze_profile(profile, client),
                "perspective2": _analyze_architecture(profile, client),
                "perspective3": _analyze_engineering(profile, client, ruler_text, classification_text),
                "perspective4b": _analyze_hidden_risks(profile, client),
            },
            run_dialectic=lambda results: _analyze_dialectic(results, client),
            run_relevance=lambda results: _score_relevance(results, user_profile, client),
            update_index=lambda request: _update_index(DEFAULT_INDEX_PATH, request),
        )

    return run_cli(pipeline, report_dir=DEFAULT_REPORT_DIR, batch_num=_next_batch_num(DEFAULT_REPORT_DIR))


if __name__ == "__main__":
    main()
